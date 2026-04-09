"""
renderers/docx_renderer.py
Generates an editable Word DOCX from CVData using python-docx.
Applies NTT Data brand styles and section formatting.
"""
from __future__ import annotations
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from renderers.pdf_renderer import _adapt

# NTT Data brand colours
BRAND_DARK = RGBColor(0x1a, 0x20, 0x35)       # Dark navy
BRAND_PRIMARY = RGBColor(0x00, 0x8B, 0x6E)    # Teal
BRAND_MUTED = RGBColor(0x5a, 0x64, 0x78)      # Grey


def _set_heading_style(paragraph, text: str, size: int = 14, color: RGBColor = BRAND_PRIMARY) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _add_section_divider(doc: Document) -> None:
    """Add a thin coloured divider line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("─" * 72)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(7)


def render_docx(cv) -> bytes:
    """
    Render a CVData object as a Word DOCX binary.
    Returns the raw bytes of the .docx file.
    """
    cv = _adapt(cv)
    doc = Document()

    # ── Page margins ──
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ── CONTACT HEADER ──
    contact = cv.contact
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(contact.full_name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = BRAND_DARK

    # Contact line
    contact_parts = []
    if contact.email:
        contact_parts.append(contact.email)
    if contact.phone:
        contact_parts.append(contact.phone)
    if contact.location:
        contact_parts.append(contact.location)
    if contact.linkedin_url:
        contact_parts.append(contact.linkedin_url)

    if contact_parts:
        contact_para = doc.add_paragraph(" · ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = BRAND_MUTED

    # ── PROFESSIONAL SUMMARY ──
    if cv.professional_summary:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "PROFESSIONAL SUMMARY", size=10)
        summary_para = doc.add_paragraph(cv.professional_summary)
        summary_para.paragraph_format.space_after = Pt(2)
        for run in summary_para.runs:
            run.font.size = Pt(10)

    # ── WORK EXPERIENCE ──
    if cv.work_experience:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "WORK EXPERIENCE", size=10)

        for role in cv.work_experience:
            # Job title + company on same line
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_before = Pt(6)
            title_run = role_para.add_run(role.job_title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = BRAND_DARK
            role_para.add_run(f"  ·  {role.company}")

            # Date + location
            date_parts = []
            if role.start_date:
                end = "Present" if role.is_current else (role.end_date or "")
                date_parts.append(f"{role.start_date} – {end}")
            if role.location:
                date_parts.append(role.location)
            if date_parts:
                date_para = doc.add_paragraph(" · ".join(date_parts))
                for run in date_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = BRAND_MUTED
                date_para.paragraph_format.space_after = Pt(3)

            # Bullet points
            bullets = role.bullet_points or (
                [role.description] if role.description else []
            )
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.25)
                bp_run = bp.add_run(bullet)
                bp_run.font.size = Pt(10)

    # ── EDUCATION ──
    if cv.education:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "EDUCATION", size=10)

        for edu in cv.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            deg_run = edu_para.add_run(edu.degree)
            deg_run.bold = True
            deg_run.font.size = Pt(11)
            edu_para.add_run(f"  ·  {edu.institution}")

            details = []
            if edu.end_date:
                details.append(edu.end_date)
            if edu.grade:
                details.append(edu.grade)
            if details:
                detail_para = doc.add_paragraph(" · ".join(details))
                for run in detail_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = BRAND_MUTED

    # ── SKILLS ──
    if cv.skills:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "SKILLS", size=10)

        # Group by category
        categories: dict[str, list[str]] = {}
        for skill in cv.skills:
            cat = skill.category or "Core Skills"
            categories.setdefault(cat, []).append(skill.name)

        for cat, skill_names in categories.items():
            skill_para = doc.add_paragraph()
            cat_run = skill_para.add_run(f"{cat}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            skills_run = skill_para.add_run(", ".join(skill_names))
            skills_run.font.size = Pt(10)

    # ── CERTIFICATIONS ──
    if cv.certifications:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "CERTIFICATIONS", size=10)

        for cert in cv.certifications:
            cert_para = doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(2)
            cert_run = cert_para.add_run(cert.name)
            cert_run.bold = True
            cert_run.font.size = Pt(10)
            if cert.issuer:
                cert_para.add_run(f"  ·  {cert.issuer}")
            if cert.date_obtained:
                cert_para.add_run(f"  ·  {cert.date_obtained}")

    # ── LANGUAGES ──
    if cv.languages:
        _add_section_divider(doc)
        h = doc.add_paragraph()
        _set_heading_style(h, "LANGUAGES", size=10)

        lang_para = doc.add_paragraph(
            "  ·  ".join(f"{l.language} ({l.proficiency})" for l in cv.languages)
        )
        for run in lang_para.runs:
            run.font.size = Pt(10)

    # ── Serialize to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
