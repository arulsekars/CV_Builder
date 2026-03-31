"""
Document Parser Agent
─────────────────────
Reads uploaded DOCX / PDF files and extracts text,
then passes to the extraction agent for structured parsing.
"""
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        text = "\n\n".join(pages)
        logger.info(f"PDF parsed: {len(text)} chars from {len(doc)} pages")
        return text
    except ImportError:
        logger.error("PyMuPDF not installed. Run: pip install pymupdf")
        raise
    except Exception as e:
        logger.error(f"PDF parse error: {e}")
        raise


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        logger.info(f"DOCX parsed: {len(text)} chars, {len(paragraphs)} paragraphs")
        return text
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        raise


def parse_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """
    Route file to correct parser based on extension.
    Returns raw text content for the extraction agent.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload PDF, DOCX, or TXT.")
